"""
zotero_field_insert.py

Insert Zotero-compatible field codes into Word (.docx) documents.

This script finds [@citationkey] markers in a .docx file and replaces them
with proper ADDIN ZOTERO_ITEM CSL_CITATION field codes that the Zotero Word
plugin recognizes. It also inserts a ZOTERO_BIBL field for the bibliography.

After running this script, open the document in Word and click "Refresh" in
the Zotero tab to renumber citations and generate the bibliography.

Usage:
    python zotero_field_insert.py input.docx [--output output.docx] [--bib references.json] [--zotero-db path]

Requirements:
    pip install python-docx

The bibliography file (--bib) should be CSL JSON with an "id" field on each
item matching the citation keys used in the document (e.g., "banwell2023").
You can export this from Zotero, or build it manually.

Optionally, if you provide a Zotero SQLite database path (--zotero-db) and
a mapping file (--keymap), the script will look up real Zotero item keys and
URIs so the plugin can link citations back to your library.
"""

import argparse
import copy
import json
import random
import re
import sqlite3
import string
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SCHEMA_URL = (
    "https://github.com/citation-style-language/schema/raw/master/csl-citation.json"
)


def random_id(length=8):
    return "".join(random.choices(string.ascii_lowercase + string.digits, k=length))


def load_bibliography(bib_path):
    """Load a CSL JSON bibliography file. Returns dict keyed by item id."""
    with open(bib_path) as f:
        items = json.load(f)
    return {item["id"]: item for item in items}


def load_keymap(keymap_path):
    """Load a JSON keymap: {cite_key: zotero_item_key_or_null}."""
    if keymap_path and Path(keymap_path).exists():
        with open(keymap_path) as f:
            return json.load(f)
    return {}


def get_zotero_item_id(zotero_db, zotero_key):
    """Look up the numeric itemID from the Zotero SQLite database."""
    if not zotero_db or not zotero_key:
        return random.randint(90000, 99999)
    try:
        db = sqlite3.connect(str(zotero_db))
        row = db.execute(
            "SELECT itemID FROM items WHERE key = ?", (zotero_key,)
        ).fetchone()
        db.close()
        return row[0] if row else random.randint(90000, 99999)
    except Exception:
        return random.randint(90000, 99999)


def get_zotero_user_id(zotero_db):
    """Try to determine the Zotero user ID from the database."""
    if not zotero_db:
        return "0"
    try:
        db = sqlite3.connect(str(zotero_db))
        row = db.execute(
            "SELECT value FROM settings WHERE setting = 'account' AND key = 'userID'"
        ).fetchone()
        db.close()
        return str(row[0]) if row else "0"
    except Exception:
        return "0"


def build_citation_json(cite_keys, bib, keymap, zotero_db, user_id):
    """Build the Zotero CSL_CITATION JSON payload for one or more keys."""
    citation_items = []
    for ck in cite_keys:
        zot_key = keymap.get(ck)
        item_id = get_zotero_item_id(zotero_db, zot_key)
        uri_key = zot_key or random_id().upper()
        uri = f"http://zotero.org/users/{user_id}/items/{uri_key}"

        csl = dict(bib.get(ck, {}))
        csl["id"] = f"{user_id}/{uri_key}"

        citation_items.append(
            {"id": item_id, "uris": [uri], "itemData": csl}
        )

    return {
        "citationID": random_id(),
        "properties": {
            "formattedCitation": "",
            "plainCitation": "",
            "noteIndex": 0,
        },
        "citationItems": citation_items,
        "schema": SCHEMA_URL,
    }


def make_superscript_rpr(font="Calibri", size_pt=11):
    """Create run properties for superscript citation text."""
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    rpr.append(fonts)
    half_pt = str(size_pt * 2)
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), half_pt)
        rpr.append(el)
    va = OxmlElement("w:vertAlign")
    va.set(qn("w:val"), "superscript")
    rpr.append(va)
    return rpr


def insert_zotero_field(parent, position, citation_json, display_text, rpr_factory):
    """Insert a 5-run Zotero ADDIN ZOTERO_ITEM field at a position in a paragraph."""
    runs = []
    for step in ("begin", "instr", "separate", "text", "end"):
        r = OxmlElement("w:r")
        r.append(copy.deepcopy(rpr_factory()))

        if step == "begin":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "begin")
            r.append(fc)
        elif step == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = (
                " ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(citation_json)
            )
            r.append(instr)
        elif step == "separate":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "separate")
            r.append(fc)
        elif step == "text":
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = display_text
            r.append(t)
        elif step == "end":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "end")
            r.append(fc)

        runs.append(r)

    for run in reversed(runs):
        parent.insert(position, run)


def insert_zotero_bibl(paragraph):
    """Replace a paragraph's content with an ADDIN ZOTERO_BIBL field."""
    for child in list(paragraph._p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink"):
            paragraph._p.remove(child)

    for step, content in [
        ("begin", None),
        (
            "instr",
            ' ADDIN ZOTERO_BIBL {"uncited":[],"omitted":[],"custom":[]} CSL_BIBLIOGRAPHY',
        ),
        ("separate", None),
        (
            "text",
            "[Bibliography will be generated by Zotero. Click Refresh in the Zotero tab.]",
        ),
        ("end", None),
    ]:
        r = OxmlElement("w:r")
        if step == "begin":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "begin")
            r.append(fc)
        elif step == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = content
            r.append(instr)
        elif step == "separate":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "separate")
            r.append(fc)
        elif step == "text":
            t = OxmlElement("w:t")
            t.text = content
            r.append(t)
        elif step == "end":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "end")
            r.append(fc)
        paragraph._p.append(r)


def process_paragraph(para, bib, keymap, zotero_db, user_id, rpr_factory):
    """Find [@key] markers in a paragraph and replace with Zotero fields."""
    full_text = para.text
    if "[@" not in full_text:
        return False

    markers = list(re.finditer(r"\[@([\w]+)\]", full_text))
    if not markers:
        return False

    p_elem = para._p

    # Collect runs and their text
    run_data = []
    for child in list(p_elem):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            texts = child.findall(qn("w:t"))
            run_text = "".join(t.text or "" for t in texts)
            rpr_elem = child.find(qn("w:rPr"))
            run_data.append(
                {
                    "text": run_text,
                    "rpr": copy.deepcopy(rpr_elem) if rpr_elem is not None else None,
                }
            )
        elif tag == "hyperlink":
            for sub_r in child.findall(qn("w:r")):
                texts = sub_r.findall(qn("w:t"))
                run_text = "".join(t.text or "" for t in texts)
                rpr_elem = sub_r.find(qn("w:rPr"))
                run_data.append(
                    {
                        "text": run_text,
                        "rpr": copy.deepcopy(rpr_elem)
                        if rpr_elem is not None
                        else None,
                    }
                )

    concat = "".join(rd["text"] for rd in run_data)
    markers = list(re.finditer(r"\[@([\w]+)\]", concat))
    if not markers:
        return False

    # Get default run properties
    default_rpr = None
    for rd in run_data:
        if rd["rpr"] is not None and rd["text"].strip():
            default_rpr = rd["rpr"]
            break

    # Clear existing runs
    for child in list(p_elem):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink"):
            p_elem.remove(child)

    # Rebuild with Zotero fields
    pos = 0
    for marker in markers:
        before = concat[pos : marker.start()]
        if before:
            r = OxmlElement("w:r")
            if default_rpr is not None:
                r.append(copy.deepcopy(default_rpr))
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = before
            r.append(t)
            p_elem.append(r)

        cite_key = marker.group(1)
        if cite_key not in bib:
            print(f"  WARNING: citation key '{cite_key}' not found in bibliography")
        cit_json = build_citation_json(
            [cite_key], bib, keymap, zotero_db, user_id
        )
        display = f"[{cite_key}]"
        insert_zotero_field(
            p_elem, len(list(p_elem)), cit_json, display, rpr_factory
        )
        pos = marker.end()

    after = concat[pos:]
    if after:
        r = OxmlElement("w:r")
        if default_rpr is not None:
            r.append(copy.deepcopy(default_rpr))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = after
        r.append(t)
        p_elem.append(r)

    return True


def main():
    parser = argparse.ArgumentParser(
        description="Insert Zotero field codes into a Word document"
    )
    parser.add_argument("input", help="Input .docx file with [@key] citation markers")
    parser.add_argument("--output", "-o", help="Output .docx path (default: overwrite input)")
    parser.add_argument("--bib", "-b", required=True, help="CSL JSON bibliography file")
    parser.add_argument("--keymap", "-k", help="JSON file mapping cite keys to Zotero item keys")
    parser.add_argument("--zotero-db", help="Path to Zotero SQLite database")
    parser.add_argument("--font", default="Calibri", help="Font for citation text (default: Calibri)")
    parser.add_argument("--size", type=int, default=11, help="Font size in pt (default: 11)")
    parser.add_argument(
        "--bib-heading",
        default="References",
        help="Heading text where bibliography should be inserted (default: References)",
    )
    args = parser.parse_args()

    output = args.output or args.input
    bib = load_bibliography(args.bib)
    keymap = load_keymap(args.keymap) if args.keymap else {}
    zotero_db = args.zotero_db
    user_id = get_zotero_user_id(zotero_db) if zotero_db else "0"

    print(f"Bibliography: {len(bib)} items")
    print(f"Keymap: {len(keymap)} entries")
    if zotero_db:
        print(f"Zotero user ID: {user_id}")

    rpr_factory = lambda: make_superscript_rpr(args.font, args.size)

    doc = Document(args.input)

    # Process citations
    n = 0
    for para in doc.paragraphs:
        if process_paragraph(para, bib, keymap, zotero_db, user_id, rpr_factory):
            n += 1
    print(f"Processed {n} paragraphs with citations")

    # Insert bibliography
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Match heading like "12. References" or just "References"
        if re.match(rf"^(\d+\.\s*)?{re.escape(args.bib_heading)}$", text):
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                # If next paragraph has text (e.g. manual refs), replace it
                insert_zotero_bibl(next_para)
            else:
                new_para = doc.add_paragraph()
                insert_zotero_bibl(new_para)
            print("Inserted ZOTERO_BIBL field")

            # Remove any remaining manual reference paragraphs
            removed = 0
            for j in range(i + 2, len(doc.paragraphs)):
                ptext = doc.paragraphs[j].text.strip()
                if re.match(r"^\d+\.\s", ptext):
                    p = doc.paragraphs[j]._p
                    parent = p.getparent()
                    if parent is not None:
                        parent.remove(p)
                        removed += 1
            if removed:
                print(f"Removed {removed} manual reference paragraphs")
            break

    doc.save(output)
    print(f"Saved: {output}")


if __name__ == "__main__":
    main()
